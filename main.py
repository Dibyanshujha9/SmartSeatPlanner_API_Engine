# main.py
# Format changes 
from fastapi import FastAPI, UploadFile, Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List
import pandas as pd
import zipfile, io, re, os
from collections import defaultdict, Counter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import PatternFill

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def insert_headers(doc, date, time, room_name):
    p1 = doc.add_paragraph(f"DATE: {date}")
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = doc.add_paragraph(f"TIME: {time}")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.runs[0].bold = True

    p3 = doc.add_paragraph(f"SEATING ARRANGEMENT FOR ROOM NO. {room_name}")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p3.runs:
        run.bold = True

def fill_columnwise(paper_queue, paper_groups, rows, cols):
    room = [["" for _ in range(cols)] for _ in range(rows)]
    dept_map = [["" for _ in range(cols)] for _ in range(rows)]
    paper_map = [["" for _ in range(cols)] for _ in range(rows)]
    seat_order = [(r, c) for c in range(cols) for r in range(rows)]
    seat_index = 0
    while seat_index < len(seat_order) and paper_queue:
        p1 = paper_queue[0]
        p2 = paper_queue[1] if len(paper_queue) > 1 else None
        for i in range(seat_index, len(seat_order)):
            r, c = seat_order[i]
            current_paper = None
            if c % 2 == 0 and paper_groups[p1]: current_paper = p1
            elif c % 2 == 1 and p2 and paper_groups[p2]: current_paper = p2
            if current_paper:
                roll, dept = paper_groups[current_paper].pop(0)
                room[r][c] = roll
                dept_map[r][c] = dept
                paper_map[r][c] = current_paper
                seat_index += 1
                if not paper_groups[current_paper]: paper_queue.remove(current_paper)
                break
            else:
                seat_index += 1
    return room, dept_map, paper_map

def dominant_dept(dept_map, col, rows):
    counts = Counter(dept_map[r][col] for r in range(rows) if dept_map[r][col])
    return counts.most_common(1)[0][0] if counts else ""

def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                tag = OxmlElement(f'w:{edge}')
                tag.set(qn('w:val'), 'single')
                tag.set(qn('w:sz'), '4')
                tag.set(qn('w:space'), '0')
                tag.set(qn('w:color'), '000000')
                borders.append(tag)
            tcPr.append(borders)

@app.post("/generate-seating-plan")
async def generate_seating_plan(
    excel_files: List[UploadFile],
    template_docx: UploadFile,
    mapping_input: str = Form(...),
    room_specs: str = Form(...),
    date: str = Form(...),
    time: str = Form(...)
):
    all_dataframes = []
    for file in excel_files:
        xls = pd.ExcelFile(await file.read())
        df_raw = xls.parse(xls.sheet_names[0], header=None)
        rows = []
        idx = 0
        while idx < len(df_raw):
            row = df_raw.iloc[idx]
            if isinstance(row[0], str) and "Paper ID" in row[0] and "Paper Code" in row[0]:
                match = re.search(r"Paper Code:\s*([A-Z0-9]+)", row[0])
                paper_code = match.group(1).strip() if match else "UNKNOWN"
                while idx < len(df_raw) and str(df_raw.iloc[idx][3]).strip().lower() != "rollno":
                    idx += 1
                idx += 1
                while idx < len(df_raw):
                    student = df_raw.iloc[idx]
                    roll = student[3]
                    name = student[5]
                    if pd.notna(roll) and isinstance(roll, (int, float)):
                        rows.append({
                            'Rollno': str(int(roll)).zfill(11),
                            'name': str(name).strip() if pd.notna(name) else "",
                            'Paper Code': paper_code
                        })
                        idx += 1
                    else:
                        break
            else:
                idx += 1
        df = pd.DataFrame(rows)
        df['last8'] = df['Rollno'].str[-8:]
        all_dataframes.append(df)

    df = pd.concat(all_dataframes, ignore_index=True)

    paper_last8_dept_map = {}
    for entry in mapping_input.split(","):
        parts = entry.strip().split("-")
        if len(parts) < 3: continue
        paper = parts[0].strip()
        dept = parts[-1].strip()
        for last8 in parts[1:-1]:
            paper_last8_dept_map[(paper, last8.strip())] = dept

    df['department'] = df.apply(lambda row: paper_last8_dept_map.get((row['Paper Code'], row['last8'])), axis=1)
    df = df[df['department'].notna()]

    parsed_rooms = []
    for spec in room_specs.split(","):
        parts = spec.strip().split(":")
        name = parts[0]
        layout = parts[2] if len(parts) == 3 else "6x8"
        rows, cols = map(int, layout.lower().split("x"))
        parsed_rooms.append((name, rows, cols))

    paper_groups = defaultdict(list)
    for _, row in df.iterrows():
        paper_groups[row['Paper Code']].append((row['Rollno'], row['department']))
    paper_sizes = {k: len(v) for k, v in paper_groups.items()}
    high = [p for p in paper_groups if paper_sizes[p] >= 10]
    low = [p for p in paper_groups if paper_sizes[p] < 10]
    paper_queue = high + low

    palette = ["F8CBAD", "DDEBF7", "C6E0B4", "F4B084", "FFD966", "D9D2E9", "B4C6E7", "E2EFDA"]
    paper_colors = {p: palette[i % len(palette)] for i, p in enumerate(paper_queue)}

    doc = Document(io.BytesIO(await template_docx.read()))
    first = True
    wb = Workbook()
    wb.remove(wb.active)

    for room_name, rows, cols in parsed_rooms:
        if not any(paper_groups.values()): break
        room, dept_map, paper_map = fill_columnwise(paper_queue, paper_groups, rows, cols)
        if not first: doc.add_page_break()
        first = False
        insert_headers(doc, date, time, room_name)

        summary = defaultdict(int)
        for r in range(rows):
            for c in range(cols):
                if room[r][c]: summary[(dept_map[r][c], paper_map[r][c])] += 1

        for (dept, paper), count in summary.items():
            para = doc.add_paragraph(f"{dept.upper()} (PAPER CODE {paper}) – {{{count}}}")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.runs[0].bold = True

        table = doc.add_table(rows=rows + 1, cols=cols)
        table.style = 'Table Grid'
        for c in range(cols):
            dpt = dominant_dept(dept_map, c, rows)
            label = "ROW-1" if c < cols // 2 else "ROW-2"
            cell = table.cell(0, c)
            cell.text = f"{dpt}\n{label}"
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if para.runs: para.runs[0].bold = True

        for r in range(rows):
            for c in range(cols):
                roll = room[r][c]
                d = dept_map[r][c]
                dom = dominant_dept(dept_map, c, rows)
                txt = roll if d == dom else f"{roll} ({d})"
                cell = table.cell(r+1, c)
                cell.text = txt
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_table_borders(table)

        sheet = wb.create_sheet(title=room_name)
        for c in range(cols):
            dpt = dominant_dept(dept_map, c, rows)
            label = "ROW-1" if c < cols // 2 else "ROW-2"
            sheet.cell(row=1, column=c+2, value=f"{dpt}\n{label}")
        for r in range(rows):
            sheet.cell(row=r+2, column=1, value=f"Row {r+1}")
            for c in range(cols):
                roll = room[r][c]
                paper = paper_map[r][c]
                cell = sheet.cell(row=r+2, column=c+2, value=roll)
                if paper:
                    color = paper_colors.get(paper, "FFFFFF")
                    cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")

    doc.save("Seating_Plan_All_Rooms.docx")
    wb.save("Seating_Summary.xlsx")
    with zipfile.ZipFile("Final_Seating_Documents.zip", "w") as zipf:
        zipf.write("Seating_Plan_All_Rooms.docx")
        zipf.write("Seating_Summary.xlsx")

    return FileResponse("Final_Seating_Documents.zip", filename="Final_Seating_Documents.zip")
